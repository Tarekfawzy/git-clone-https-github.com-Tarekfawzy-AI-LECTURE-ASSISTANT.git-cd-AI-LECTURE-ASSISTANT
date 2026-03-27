from io import BytesIO
from typing import Union
import io

def export_to_txt(content: str) -> bytes:
    """Export content to TXT format."""
    return content.encode('utf-8')

def export_to_pdf(content: str, title: str = "Document") -> bytes:
    """Export content to PDF format."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#1f77b4',
            spaceAfter=30
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Add content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph, styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    except ImportError:
        # Fallback to simple text if reportlab not available
        return content.encode('utf-8')

def export_to_docx(content: str, title: str = "Document") -> bytes:
    """Export content to DOCX format."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph)
                p.style = 'Normal'
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    except ImportError:
        # Fallback to simple text if python-docx not available
        return content.encode('utf-8')
